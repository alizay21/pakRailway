import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_paragraph_with_spacing(doc, text="", style=None, before=0, after=6, line_spacing=1.15):
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    return p

def main():
    doc = Document()
    
    # 1. Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Enable different first page for title page
        section.different_first_page_header_footer = True
        
        # Configure headers and footers for subsequent pages
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "Software Project Management (SPM) PBL Report  |  PakRail"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.style.font.name = 'Calibri'
        hp.style.font.size = Pt(8.5)
        hp.style.font.color.rgb = RGBColor(120, 120, 120)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "BSSE-6, Riphah International University  |  Page "
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.style.font.name = 'Calibri'
        fp.style.font.size = Pt(8.5)
        fp.style.font.color.rgb = RGBColor(120, 120, 120)
        # Add dynamic page field if possible, or leave clean spacing
        f_run = fp.add_run("[Dynamic Page Number]")
        f_run.italic = True

    # Palette
    COLOR_PRIMARY = RGBColor(27, 54, 93)     # Deep Navy Blue (#1B365D)
    COLOR_SECONDARY = RGBColor(15, 118, 110) # Emerald Green (#0F766E)
    COLOR_TEXT = RGBColor(51, 51, 51)        # Dark Charcoal
    HEX_PRIMARY = "1B365D"
    HEX_SECONDARY = "0F766E"
    HEX_LIGHT_GREY = "F3F4F6"
    HEX_BORDER_GREY = "D1D5DB"

    # Set default font
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT

    # Set heading styles
    for h_name, size, color, bold in [('Heading 1', 18, COLOR_PRIMARY, True), 
                                      ('Heading 2', 14, COLOR_SECONDARY, True), 
                                      ('Heading 3', 12, COLOR_PRIMARY, True)]:
        h_style = doc.styles[h_name]
        h_style.font.name = 'Calibri'
        h_style.font.size = Pt(size)
        h_style.font.color.rgb = color
        h_style.font.bold = bold

    # ==================== TITLE PAGE ====================
    # Spacing before title
    for _ in range(3):
        add_paragraph_with_spacing(doc)
        
    p_pbl = add_paragraph_with_spacing(doc, "PROJECT-BASED LEARNING (PBL) REPORT", before=12, after=6)
    p_pbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pbl.runs[0].font.size = Pt(14)
    p_pbl.runs[0].font.bold = True
    p_pbl.runs[0].font.color.rgb = COLOR_SECONDARY
    
    p_title = add_paragraph_with_spacing(doc, "PakRail – Smart Railway Reservation & Seat Management System", before=6, after=18)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.runs[0].font.size = Pt(24)
    p_title.runs[0].font.bold = True
    p_title.runs[0].font.color.rgb = COLOR_PRIMARY
    
    p_course = add_paragraph_with_spacing(doc, "Course: Software Project Management (SPM)\nClass: BSSE-6", before=12, after=24)
    p_course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_course.runs[0].font.size = Pt(12)
    p_course.runs[0].font.italic = True
    
    for _ in range(3):
        add_paragraph_with_spacing(doc)
        
    # Student Metadata Table
    p_team_title = add_paragraph_with_spacing(doc, "Project Team Members", before=12, after=6)
    p_team_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_team_title.runs[0].font.bold = True
    p_team_title.runs[0].font.size = Pt(12)
    
    table_team = doc.add_table(rows=3, cols=2)
    table_team.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_team.rows[0].cells
    hdr_cells[0].text = 'Member Name'
    hdr_cells[1].text = 'SAP ID'
    for cell in hdr_cells:
        set_cell_background(cell, HEX_PRIMARY)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        
    m1_cells = table_team.rows[1].cells
    m1_cells[0].text = 'Alizay Aman Niazi (Team Leader / PM)'
    m1_cells[1].text = '58604'
    m2_cells = table_team.rows[2].cells
    m2_cells[0].text = 'Ayesha Kanwal (Technical Analyst)'
    m2_cells[1].text = '57505'
    
    for row in table_team.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            set_cell_background(cell, HEX_LIGHT_GREY)

    add_paragraph_with_spacing(doc, "\n")
    
    p_inst = add_paragraph_with_spacing(doc, "Instructor:\nSir Waqar / Dr. Waqar Mehmood", before=12, after=24)
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.runs[0].font.bold = True
    p_inst.runs[0].font.size = Pt(11)
    
    p_univ = add_paragraph_with_spacing(doc, "Faculty of Computing\nRiphah International University, Islamabad", before=12, after=12)
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_univ.runs[0].font.size = Pt(12)
    p_univ.runs[0].font.bold = True
    
    p_tech = add_paragraph_with_spacing(doc, "Technology Stack:\nReact.js • JavaScript • Vite • Tailwind CSS • Node.js • Express.js • MongoDB", before=18, after=12)
    p_tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tech.runs[0].font.size = Pt(9.5)
    p_tech.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_page_break()

    # ==================== MAIN REPORT CONTENT ====================
    
    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    p = add_paragraph_with_spacing(doc, 
        "Modern railway systems form the backbone of national transportation infrastructure, handling millions of passengers daily. "
        "However, legacy reservation setups frequently struggle with operational bottlenecks including long ticket counter queues, manual "
        "booking transcription errors, seat double-booking, lack of real-time seat inventory visibility, poor reporting analytics, "
        "and limited accessibility. These challenges lead to passenger dissatisfaction and revenue leakage for railway administrations.", 
        before=6, after=6)
    p = add_paragraph_with_spacing(doc, 
        "To resolve these challenges, the PakRail system is introduced as a digital transformation initiative. PakRail is a web-based "
        "smart railway reservation and seat management system designed to automate ticketing operations. It provides passengers with "
        "the ability to search trains by route and date, view real-time visual seat maps, book specific seats, enter passenger details, "
        "simulate credit card/cash payments, track bookings using unique Passenger Name Records (PNR), download ticket receipts in PDF, "
        "and request cancellations. Concurrently, it offers administrators a robust dashboard equipped with real-time statistics, "
        "train scheduling interfaces, comprehensive booking reports, and visual charts depicting operational performance.", 
        before=6, after=6)
    p = add_paragraph_with_spacing(doc, 
        "This project has been developed as a Software Project Management (SPM) Project-Based Learning (PBL) assignment. It employs the "
        "MERN (MongoDB, Express.js, React, Node.js) technology stack, adhering to professional software engineering standards, "
        "project scheduling, resource allocation, and risk management strategies to deliver a reliable, secure, and user-centric platform.", 
        before=6, after=12)

    # 2. Introduction
    doc.add_heading("2. Introduction", level=1)
    p = add_paragraph_with_spacing(doc, 
        "The railway ticketing domain represents a high-concurrency, transaction-critical system requiring high data integrity. "
        "Digital railway systems are essential for optimizing seat inventory utilization, preventing double bookings, and providing "
        "frictionless service access to remote passengers. Traditional paper-based ticket booking systems fail to scale, creating "
        "administrative overhead and restricting passenger autonomy.", 
        before=6, after=6)
    p = add_paragraph_with_spacing(doc, 
        "PakRail offers a web-based portal serving two primary classes of users: Passengers and Administrators. Passengers can plan "
        "journeys, select preferred seats, and finalize bookings from their own devices. Administrators can manage the master schedule "
        "of active trains, view system-wide seat configurations, audit financial bookings, and analyze passenger booking trends.", 
        before=6, after=6)
    p = add_paragraph_with_spacing(doc, 
        "The system has been configured and deployed locally for validation. The user interface can be accessed at: http://localhost:5173/. "
        "The backend API runs on port 5000, communicating with a local MongoDB database to maintain real-time status synchronization.", 
        before=6, after=12)

    # 3. Problem Statement
    doc.add_heading("3. Problem Statement", level=1)
    p = add_paragraph_with_spacing(doc, 
        "Legacy passenger transportation networks in Pakistan rely on manual, centralized booking offices. This outdated framework "
        "imposes significant delays, limits capacity management, and introduces critical operational vulnerabilities. The core problems "
        "addressed by the PakRail system include:", 
        before=6, after=6)

    doc.add_heading("3.1 Manual Ticket Booking Delays", level=2)
    add_paragraph_with_spacing(doc, 
        "Passengers must physically visit train stations, standing in long queues for ticket purchases. This creates delays and "
        "unnecessary transit costs, particularly during holidays and peak travel seasons.", after=6)

    doc.add_heading("3.2 Lack of Real-Time Seat Availability", level=2)
    add_paragraph_with_spacing(doc, 
        "There is no mechanism for passengers to check available seats prior to visiting the ticket window, resulting in speculative "
        "travel planning and wasted trips when trains are fully occupied.", after=6)

    doc.add_heading("3.3 Overbooking and Seat Allocation Errors", level=2)
    add_paragraph_with_spacing(doc, 
        "In manual bookkeeping systems, seat allocation data is prone to replication lags. This leads to double-booking the same physical "
        "seat for multiple passengers on overlapping journeys, creating passenger conflicts.", after=6)

    doc.add_heading("3.4 Inefficient Booking Record Management", level=2)
    add_paragraph_with_spacing(doc, 
        "Managing ledger books and physical receipts makes it challenging to search, update, or cancel bookings. Retelling lost receipts "
        "requires tedious manual cross-referencing, exposing the railway to operational inefficiencies.", after=6)

    doc.add_heading("3.5 Limited Accessibility for Passengers", level=2)
    add_paragraph_with_spacing(doc, 
        "Passengers cannot cancel bookings remotely, nor can they retrieve booking details online using digital identifiers. Language "
        "barriers also exist as most booking interfaces are English-only, neglecting Urdu-only speakers.", after=6)

    doc.add_heading("3.6 Weak Reporting and Monitoring for Railway Management", level=2)
    add_paragraph_with_spacing(doc, 
        "Station masters and railway administrators lack access to consolidated reports regarding seat load factors, class-specific revenues, "
        "and passenger counts, which inhibits data-driven scheduling and tariff adjustments.", after=6)

    doc.add_heading("3.7 Need for a Centralized Digital Reservation System", level=2)
    add_paragraph_with_spacing(doc, 
        "Without a unified, centralized transactional database, scheduling trains, managing seat pricing classes, and generating "
        "financial audits requires extensive physical paperwork, increasing administrative costs and the potential for database fraud.", 
        after=12)

    doc.add_page_break()

    # 4. Project Objectives
    doc.add_heading("4. Project Objectives", level=1)
    p = add_paragraph_with_spacing(doc, "The main objectives of the PakRail project are:", before=6, after=6)
    
    objectives = [
        "To automate the railway ticketing process, allowing remote ticket reservation and cancellation.",
        "To enable passengers to search for trains by departure/destination cities and travel dates.",
        "To provide a real-time visual seat map showing available, selected, and booked seats for precise selection.",
        "To prevent double-booking issues by locking seats during transaction processing via atomic database updates.",
        "To generate unique Passenger Name Records (PNR) for secure tracking, validation, and status checking.",
        "To offer simulated checkout options, including simulated credit card payment and a Pay-at-Counter reservation hold.",
        "To generate downloadable PDF tickets containing journey details, passenger information, seat numbers, and PNR codes.",
        "To build an administrative panel for scheduling trains, auditing active bookings, and analyzing business performance.",
        "To integrate bilingual support (English and Urdu) with automated Right-to-Left (RTL) layout switching for accessibility.",
        "To improve overall railway operational efficiency, record transparency, and passenger booking convenience."
    ]
    for i, obj in enumerate(objectives, 1):
        p_list = add_paragraph_with_spacing(doc, f"{i}. {obj}", before=0, after=4)
        p_list.paragraph_format.left_indent = Inches(0.25)
        
    add_paragraph_with_spacing(doc, "", after=12)

    # 5. Project Scope
    doc.add_heading("5. Project Scope", level=1)
    p = add_paragraph_with_spacing(doc, 
        "The PakRail project is categorized into functional modules defining user features, administrative capabilities, "
        "and system boundaries.", before=6, after=6)

    doc.add_heading("Passenger Module", level=2)
    p_scoped = [
        "Registration & Login: Secure passenger account creation and validation using email and passwords.",
        "Train Search: Query active trains based on departure city, arrival city, travel class, and date.",
        "Real-Time Seat Map: A grid layout representing seats inside the train coach showing real-time reservation statuses.",
        "Seat Selection: Allows passengers to choose up to 6 seats per transaction.",
        "Passenger Details Form: Capture passenger names, ages, genders, and 13-digit CNIC numbers.",
        "Simulated Payments: Interactive form supporting simulated credit card transactions or a Pay-at-Counter booking hold.",
        "PNR Tracking: Search page to look up ticket confirmation details, seat lists, and payment status using a PNR code.",
        "PDF Ticket Download: Generates a printable receipt using the browser-side jsPDF script.",
        "Booking History & Cancellation: A dashboard showing past trips with options to request booking cancellations."
    ]
    for bullet in p_scoped:
        p_b = add_paragraph_with_spacing(doc, f"• {bullet}", before=0, after=4)
        p_b.paragraph_format.left_indent = Inches(0.25)

    doc.add_heading("Admin Module", level=2)
    a_scoped = [
        "Admin Dashboard: High-level dashboard showing total revenue, active trains, booked seats, and user registrations.",
        "Train Management: Create, update, view, and delete train routes, schedules, seat capacities, and base fares.",
        "Booking Audits: View, filter, search, and cancel passenger bookings directly from the administration interface.",
        "Reports & Analytics: Dynamic charts (powered by Recharts) showing booking volumes and revenue distributions."
    ]
    for bullet in a_scoped:
        p_b = add_paragraph_with_spacing(doc, f"• {bullet}", before=0, after=4)
        p_b.paragraph_format.left_indent = Inches(0.25)

    doc.add_heading("System Module", level=2)
    s_scoped = [
        "JWT Authentication: Session management and security tokens for client-server communication.",
        "MongoDB Schema Mappings: Mongoose models handling relations between Users, Trains, Bookings, and Seats.",
        "Bilingual Localization: Frontend translations utilizing i18next with RTL layouts configured for Urdu.",
        "Atomic Seat Locks: Transacted updates ensuring seat statuses are locked when a booking is created."
    ]
    for bullet in s_scoped:
        p_b = add_paragraph_with_spacing(doc, f"• {bullet}", before=0, after=4)
        p_b.paragraph_format.left_indent = Inches(0.25)

    doc.add_heading("Out of Scope", level=2)
    out_scoped = [
        "Integration with Pakistan Railways official servers or live operational APIs.",
        "Real-world commercial bank credit card checkout processing (such as Stripe or Payoneer integration).",
        "GPS-based live train location tracking via mobile telemetry or satellite positioning.",
        "SMS gateway notifications (replaced by UI success messages and PNR lookups due to cost constraints)."
    ]
    for bullet in out_scoped:
        p_b = add_paragraph_with_spacing(doc, f"• {bullet}", before=0, after=4)
        p_b.paragraph_format.left_indent = Inches(0.25)

    add_paragraph_with_spacing(doc, "", after=12)

    doc.add_page_break()

    # 6. Stakeholder Analysis
    doc.add_heading("6. Stakeholder Analysis", level=1)
    p = add_paragraph_with_spacing(doc, 
        "Stakeholders are individuals or entities with an interest in the project's development or deployment. "
        "They are divided into primary, secondary, and external categories.", before=6, after=6)

    # 6.1 Primary
    doc.add_heading("6.1 Primary Stakeholders", level=2)
    table_sh1 = doc.add_table(rows=4, cols=4)
    table_sh1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Stakeholder', 'Interest', 'Role', 'Influence Level']
    for idx, name in enumerate(headers):
        table_sh1.rows[0].cells[idx].text = name
        set_cell_background(table_sh1.rows[0].cells[idx], HEX_PRIMARY)
        table_sh1.rows[0].cells[idx].paragraphs[0].runs[0].font.bold = True
        table_sh1.rows[0].cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(table_sh1.rows[0].cells[idx], top=120, bottom=120, left=100, right=100)

    sh1_data = [
        ['Passengers', 'Convenient booking, seat choices, tracking.', 'End users booking tickets.', 'High'],
        ['Railway Admin', 'Efficient seat allocation, schedule management.', 'Manage trains, routes, and bookings.', 'High'],
        ['Project Team', 'Academic success, software testing, deployment.', 'Developers, managers, documenters.', 'Medium']
    ]
    for r_idx, row in enumerate(sh1_data, 1):
        for c_idx, val in enumerate(row):
            cell = table_sh1.rows[r_idx].cells[c_idx]
            cell.text = val
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            if r_idx % 2 == 0:
                set_cell_background(cell, HEX_LIGHT_GREY)

    # 6.2 Secondary
    doc.add_heading("6.2 Secondary Stakeholders", level=2)
    table_sh2 = doc.add_table(rows=4, cols=4)
    table_sh2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, name in enumerate(headers):
        table_sh2.rows[0].cells[idx].text = name
        set_cell_background(table_sh2.rows[0].cells[idx], HEX_SECONDARY)
        table_sh2.rows[0].cells[idx].paragraphs[0].runs[0].font.bold = True
        table_sh2.rows[0].cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(table_sh2.rows[0].cells[idx], top=120, bottom=120, left=100, right=100)

    sh2_data = [
        ['Ticket Counter Staff', 'Reduced manual workload, off-counter queues.', 'Assist physical ticket booking.', 'Medium'],
        ['QA/Testers', 'Ensuring bug-free and robust seat mapping.', 'Validate features and system routes.', 'Medium'],
        ['Academic Supervisor', 'Academic grading, SPM standards alignment.', 'Guidance, grading, progress audits.', 'High']
    ]
    for r_idx, row in enumerate(sh2_data, 1):
        for c_idx, val in enumerate(row):
            cell = table_sh2.rows[r_idx].cells[c_idx]
            cell.text = val
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            if r_idx % 2 == 0:
                set_cell_background(cell, HEX_LIGHT_GREY)

    # 6.3 External
    doc.add_heading("6.3 External Stakeholders", level=2)
    table_sh3 = doc.add_table(rows=4, cols=4)
    table_sh3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, name in enumerate(headers):
        table_sh3.rows[0].cells[idx].text = name
        set_cell_background(table_sh3.rows[0].cells[idx], HEX_PRIMARY)
        table_sh3.rows[0].cells[idx].paragraphs[0].runs[0].font.bold = True
        table_sh3.rows[0].cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(table_sh3.rows[0].cells[idx], top=120, bottom=120, left=100, right=100)

    sh3_data = [
        ['Payment Providers', 'Future real payment gateway merchant.', 'Process real card payments.', 'Low'],
        ['Transport Authority', 'Regulatory compliance, safety audits.', 'Legal oversight of operations.', 'Medium'],
        ['Hosting Providers', 'Server uptime, deployment resources.', 'Provide cloud instances (e.g. AWS).', 'Low']
    ]
    for r_idx, row in enumerate(sh3_data, 1):
        for c_idx, val in enumerate(row):
            cell = table_sh3.rows[r_idx].cells[c_idx]
            cell.text = val
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            if r_idx % 2 == 0:
                set_cell_background(cell, HEX_LIGHT_GREY)

    add_paragraph_with_spacing(doc, "", after=12)

    doc.add_page_break()

    # 7. Feasibility Study
    doc.add_heading("7. Feasibility Study", level=1)
    p = add_paragraph_with_spacing(doc, 
        "A feasibility study assesses whether the proposed PakRail project can be developed given constraints of "
        "technology, cost, operations, and scheduling.", before=6, after=6)

    doc.add_heading("7.1 Technical Feasibility", level=2)
    p = add_paragraph_with_spacing(doc, 
        "The project is technically feasible. The team is utilizing the MERN stack, which provides several technical advantages:\n"
        "• React.js & Vite: Fast, responsive client rendering with Hot Module Replacement (HMR).\n"
        "• Tailwind CSS: Rapid styled UI creation without heavy custom CSS writing.\n"
        "• Node.js & Express: Lightweight, high-throughput backend API routing.\n"
        "• MongoDB & Mongoose: Flexible document schema handling, allowing seat maps to be queried efficiently.\n"
        "• jsPDF & Recharts: Out-of-the-box libraries for ticket printing and dashboard statistics respectively.", 
        before=6, after=6)

    doc.add_heading("7.2 Economic Feasibility", level=2)
    p = add_paragraph_with_spacing(doc, 
        "For an academic project, the economic footprint is minimal because developer resources (laptops) and tools "
        "(VS Code, MongoDB Community Server) are open-source and already owned. The cost table below outlines "
        "the academic development budget in PKR:", before=6, after=6)
        
    table_cost = doc.add_table(rows=8, cols=2)
    table_cost.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_cost.rows[0].cells[0].text = 'Resource / Expense Category'
    table_cost.rows[0].cells[1].text = 'Cost (PKR)'
    for cell in table_cost.rows[0].cells:
        set_cell_background(cell, HEX_PRIMARY)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        
    cost_data = [
        ['Development Laptops (Academic Owned)', 'Rs. 0'],
        ['Internet Charges (5 Months)', 'Rs. 8,000'],
        ['Project Documentation and Report Printing', 'Rs. 3,000'],
        ['Demo Web Hosting (Render/Vercel/MongoDB Atlas)', 'Rs. 5,000'],
        ['Development Tools & IDEs (VS Code, Postman - Free Tier)', 'Rs. 0'],
        ['Testing Resources (Simulated user validation)', 'Rs. 2,000'],
        ['Miscellaneous & Contingency', 'Rs. 2,000']
    ]
    for idx, row in enumerate(cost_data, 1):
        table_cost.rows[idx].cells[0].text = row[0]
        table_cost.rows[idx].cells[1].text = row[1]
        for cell in table_cost.rows[idx].cells:
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if idx % 2 == 0:
                set_cell_background(cell, HEX_LIGHT_GREY)
                
    # Add Total Row
    total_row = table_cost.add_row()
    total_row.cells[0].text = 'Total Estimated Academic Cost'
    total_row.cells[1].text = 'Rs. 20,000'
    for cell in total_row.cells:
        set_cell_background(cell, HEX_SECONDARY)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

    p = add_paragraph_with_spacing(doc, 
        "Note: In a commercial deployment, hosting, continuous maintenance, paid SMTP/SMS servers, and payment "
        "merchant licenses would increase the budget significantly.", before=6, after=6)

    doc.add_heading("7.3 Operational Feasibility", level=2)
    p = add_paragraph_with_spacing(doc, 
        "The system exhibits high operational feasibility. The interfaces are highly intuitive, and the passenger booking path "
        "is direct (Search → Select Seats → Enter Details → simulated Checkout → Ticket Download). The admin dashboard provides "
        "one-click management over trains and bookings. Little to no training is required for basic operations.", before=6, after=6)

    doc.add_heading("7.4 Schedule Feasibility", level=2)
    p = add_paragraph_with_spacing(doc, 
        "The project is structured to fit within the 16-week academic semester. By employing Agile Scrum methodologies, "
        "the team has divided development tasks into 2-week sprints, ensuring schedule compliance.", before=6, after=12)

    doc.add_page_break()

    # 8. System Architecture
    doc.add_heading("8. System Architecture", level=1)
    p = add_paragraph_with_spacing(doc, 
        "PakRail utilizes a classic 3-Tier Web Architecture which isolates client representation, API logic, and data storage layers.", 
        before=6, after=6)
    
    # Textual diagram description
    p_diag = add_paragraph_with_spacing(doc, 
        "Passenger/Admin Browser  <--->  React Frontend (Vite)  <--->  Express API Server (Node.js)  <--->  MongoDB Database", 
        before=12, after=12)
    p_diag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_diag.runs[0].font.bold = True
    p_diag.runs[0].font.size = Pt(11.5)
    p_diag.runs[0].font.color.rgb = COLOR_SECONDARY
    
    p = add_paragraph_with_spacing(doc, 
        "• Frontend Presentation Layer: Developed in React.js. It handles user authentication contexts, seat-selection maps, "
        "and client-side state transitions. It communicates with the backend via Axios HTTP requests.\n"
        "• Backend Application Layer: Designed in Node.js with the Express.js framework. It acts as a RESTful controller, "
        "verifying JWT signatures, performing business validations (such as seat occupancy conflicts), and exposing CRUD API routes.\n"
        "• Database Storage Layer: Powered by MongoDB. Mongoose is utilized to enforce schema validation rules on the unstructured "
        "NoSQL database, ensuring correct relationship mapping between trains, bookings, users, and seats.", 
        before=6, after=12)

    # 9. Resource Planning
    doc.add_heading("9. Resource Planning", level=1)
    p = add_paragraph_with_spacing(doc, "Resource planning ensures adequate allocation of human, hardware, and software assets.", before=6, after=6)
    
    doc.add_heading("9.1 Human Resources", level=2)
    table_hr = doc.add_table(rows=3, cols=3)
    table_hr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr_headers = ['Name', 'Project Role', 'Responsibilities']
    for idx, name in enumerate(hr_headers):
        table_hr.rows[0].cells[idx].text = name
        set_cell_background(table_hr.rows[0].cells[idx], HEX_PRIMARY)
        table_hr.rows[0].cells[idx].paragraphs[0].runs[0].font.bold = True
        table_hr.rows[0].cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(table_hr.rows[0].cells[idx], top=120, bottom=120, left=150, right=150)
        
    hr_data = [
        ['Alizay Aman Niazi', 'Project Manager / Lead Developer', 'Database design, backend routing, security controllers, schedule tracking.'],
        ['Ayesha Kanwal', 'Technical Analyst / Documentation Coordinator', 'Requirement gathering, UI components coordination, testing case creation, final report compilation.']
    ]
    for idx, row in enumerate(hr_data, 1):
        table_hr.rows[idx].cells[0].text = row[0]
        table_hr.rows[idx].cells[1].text = row[1]
        table_hr.rows[idx].cells[2].text = row[2]
        for cell in table_hr.rows[idx].cells:
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if idx % 2 == 0:
                set_cell_background(cell, HEX_LIGHT_GREY)

    doc.add_heading("9.2 Hardware Resources", level=2)
    p = add_paragraph_with_spacing(doc, 
        "• Development Workstations: 2 Core-i7 Laptops with 16GB RAM and 512GB SSDs.\n"
        "• Network Connection: 20 Mbps broadband internet for packages installation and remote database syncing.\n"
        "• Testing Devices: Android & iOS smartphones to test responsiveness of tailwind layout classes.", 
        before=6, after=6)

    doc.add_heading("9.3 Software Resources", level=2)
    p = add_paragraph_with_spacing(doc, 
        "• IDE & Tools: Visual Studio Code, Postman API client, and Git / GitHub.\n"
        "• DB Engine: MongoDB Community Server (v8.2) & MongoDB Compass UI.\n"
        "• Browsers: Google Chrome & Microsoft Edge for validation and testing.\n"
        "• Report Tools: Microsoft Word, Microsoft Excel, Draw.io.", 
        before=6, after=12)

    doc.add_page_break()

    # 10. Project Schedule
    doc.add_heading("10. Project Schedule", level=1)
    p = add_paragraph_with_spacing(doc, 
        "The PakRail project timeline spanned 16 weeks, tracking the standard software development lifecycle (SDLC) phases:", 
        before=6, after=6)
        
    table_sched = doc.add_table(rows=9, cols=3)
    table_sched.alignment = WD_TABLE_ALIGNMENT.CENTER
    sc_headers = ['Phase', 'Duration', 'Activities Involved']
    for idx, name in enumerate(sc_headers):
        table_sched.rows[0].cells[idx].text = name
        set_cell_background(table_sched.rows[0].cells[idx], HEX_PRIMARY)
        table_sched.rows[0].cells[idx].paragraphs[0].runs[0].font.bold = True
        table_sched.rows[0].cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(table_sched.rows[0].cells[idx], top=120, bottom=120, left=150, right=150)
        
    sched_data = [
        ['Project Initiation', 'Weeks 1–2', 'Ideation, team registration, proposal submission.'],
        ['Project Planning', 'Weeks 3–5', 'Scope determination, feasibility study, scheduling, cost assessment.'],
        ['Requirement Analysis', 'Weeks 5–6', 'System specifications gathering, functional & non-functional scoping.'],
        ['System Design', 'Weeks 6–7', 'Mongoose database model design, UI layout drafting, routes plotting.'],
        ['Development & Coding', 'Weeks 8–12', 'JWT auth build, backend Express controllers, React frontend views, seat grid binding.'],
        ['Integration & Testing', 'Weeks 12–14', 'Postman route audits, manual passenger booking path validations, security checkups.'],
        ['Localhost Demo & Launch', 'Weeks 15', 'Local database seeding, frontend build compilation, localhost verification.'],
        ['Final Documentation', 'Weeks 16–17', 'Report compilation, user validation, submission and project defense.']
    ]
    for idx, row in enumerate(sched_data, 1):
        table_sched.rows[idx].cells[0].text = row[0]
        table_sched.rows[idx].cells[1].text = row[1]
        table_sched.rows[idx].cells[2].text = row[2]
        for cell in table_sched.rows[idx].cells:
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if idx % 2 == 0:
                set_cell_background(cell, HEX_LIGHT_GREY)

    add_paragraph_with_spacing(doc, "", after=12)

    doc.add_page_break()

    # 11. Risk Management
    doc.add_heading("11. Risk Management", level=1)
    p = add_paragraph_with_spacing(doc, 
        "Risk management identifies, logs, and defines mitigations for technical and organizational issues that may "
        "arise during the project lifecycle.", before=6, after=6)
        
    table_risk = doc.add_table(rows=11, cols=4)
    table_risk.alignment = WD_TABLE_ALIGNMENT.CENTER
    rk_headers = ['Identified Risk', 'Probability', 'Impact', 'Mitigation Strategy']
    for idx, name in enumerate(rk_headers):
        table_risk.rows[0].cells[idx].text = name
        set_cell_background(table_risk.rows[0].cells[idx], HEX_PRIMARY)
        table_risk.rows[0].cells[idx].paragraphs[0].runs[0].font.bold = True
        table_risk.rows[0].cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(table_risk.rows[0].cells[idx], top=120, bottom=120, left=100, right=100)
        
    risk_data = [
        ['Requirement Creep', 'Medium', 'Medium', 'Lock scope early; implement backlog items only in future versions.'],
        ['Resource Unavailability', 'Low', 'High', 'Cross-train team members on both frontend and backend tasks.'],
        ['Development Slippage', 'Medium', 'Medium', 'Utilize 2-week milestones; maintain tasks checklist on GitHub.'],
        ['MongoDB Connection Issues', 'Medium', 'High', 'Configure static database fallback (local mongodb service).'],
        ['Seat Double-Booking Bug', 'Low', 'High', 'Implement atomic seat status locks on backend booking API.'],
        ['JWT Token Exploit', 'Low', 'Medium', 'Enforce secure token signing keys and set appropriate expiration times.'],
        ['Localhost Demo Failure', 'Low', 'High', 'Maintain offline seed files (seed.js) for quick DB population.'],
        ['Bilingual Layout Overlaps', 'Medium', 'Low', 'Validate RTL alignment styling using flexible Flexbox layouts.'],
        ['Incomplete Documentation', 'Low', 'Medium', 'Compile documentation in parallel with coding sprint reviews.'],
        ['Performance Sluggishness', 'Low', 'Medium', 'Optimize database queries via selective schema indexing.']
    ]
    for idx, row in enumerate(risk_data, 1):
        table_risk.rows[idx].cells[0].text = row[0]
        table_risk.rows[idx].cells[1].text = row[1]
        table_risk.rows[idx].cells[2].text = row[2]
        table_risk.rows[idx].cells[3].text = row[3]
        for cell in table_risk.rows[idx].cells:
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            if idx % 2 == 0:
                set_cell_background(cell, HEX_LIGHT_GREY)

    add_paragraph_with_spacing(doc, "", after=12)

    doc.add_page_break()

    # 12. Testing Strategy
    doc.add_heading("12. Testing Strategy", level=1)
    p = add_paragraph_with_spacing(doc, 
        "A multi-tiered testing methodology was executed to guarantee functional completeness, data integrity, "
        "and security compliance.", before=6, after=6)
        
    p_test = add_paragraph_with_spacing(doc, 
        "• Functional Testing: Evaluated registration, search filters, seat selection controls, details submission, payment, and cancellation.\n"
        "• Integration Testing: Audited communication between React Axios requests and Express API routes.\n"
        "• Database Testing: Verified that Seat collections status toggle correctly and that user and booking details persist.\n"
        "• Security Testing: Audited bcrypt hashing and JWT protected route redirects on passenger-only pages.\n"
        "• Usability Testing: Verified that booking could be completed within 2 minutes without user interface confusion.\n"
        "• Performance Testing: Checked backend API query response times for seat status retrieval (under 100ms on localhost).", 
        before=6, after=6)

    doc.add_heading("12.1 Sample Test Cases Table", level=2)
    table_tc = doc.add_table(rows=6, cols=5)
    table_tc.alignment = WD_TABLE_ALIGNMENT.CENTER
    tc_headers = ['Test ID', 'Module Name', 'Test Scenario Description', 'Expected Result', 'Status']
    for idx, name in enumerate(tc_headers):
        table_tc.rows[0].cells[idx].text = name
        set_cell_background(table_tc.rows[0].cells[idx], HEX_SECONDARY)
        table_tc.rows[0].cells[idx].paragraphs[0].runs[0].font.bold = True
        table_tc.rows[0].cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(table_tc.rows[0].cells[idx], top=120, bottom=120, left=100, right=100)
        
    tc_data = [
        ['TC_AUTH_01', 'Auth Module', 'Register user with missing phone field.', 'Registration fails showing validation error.', 'Passed'],
        ['TC_SRCH_02', 'Search Module', 'Query Lahore to Karachi for past travel date.', 'Search button disabled / past date error shown.', 'Passed'],
        ['TC_SEAT_03', 'Seat Map', 'Select A1 and proceed to checkout.', 'Seat locks successfully; proceeding allowed.', 'Passed'],
        ['TC_PAY_04', 'Payment', 'Select Cash and complete reservation hold.', 'Booking confirmed under Pending status with PNR.', 'Passed'],
        ['TC_ADM_05', 'Admin Module', 'Access /admin route as a basic passenger.', 'Redirected to home page automatically.', 'Passed']
    ]
    for idx, row in enumerate(tc_data, 1):
        table_tc.rows[idx].cells[0].text = row[0]
        table_tc.rows[idx].cells[1].text = row[1]
        table_tc.rows[idx].cells[2].text = row[2]
        table_tc.rows[idx].cells[3].text = row[3]
        table_tc.rows[idx].cells[4].text = row[4]
        for cell in table_tc.rows[idx].cells:
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            if idx % 2 == 0:
                set_cell_background(cell, HEX_LIGHT_GREY)

    add_paragraph_with_spacing(doc, "", after=12)

    doc.add_page_break()

    # 13. Expected Benefits
    doc.add_heading("13. Expected Benefits", level=1)
    p = add_paragraph_with_spacing(doc, 
        "Implementing PakRail delivers direct values to operations and indirect benefits to the railway ecosystem.", 
        before=6, after=6)

    doc.add_heading("13.1 Direct Benefits", level=2)
    p = add_paragraph_with_spacing(doc, 
        "• Load reduction: Offloads 60%+ of physical ticket office queue volumes to remote channels.\n"
        "• Error Reduction: Eliminate manual seat write mistakes and duplication conflicts via atomic seat allocation.\n"
        "• Fast Processing: Simplifies the cancellation workflow, removing the need for manual record audits.", 
        before=6, after=6)

    doc.add_heading("13.2 Indirect Benefits", level=2)
    p = add_paragraph_with_spacing(doc, 
        "• Elevated Satisfaction: Passengers gain control over seat preferences (Economy vs. First Class).\n"
        "• Data Transparency: Passengers can verify seats status in real-time without administrative intermediaries.\n"
        "• Analytical Control: Enables management to make schedule updates backed by revenue and boarding reports.", 
        before=6, after=12)

    doc.add_page_break()

    # 14. Screenshots of Implemented System
    doc.add_heading("14. Screenshots of Implemented System", level=1)
    p = add_paragraph_with_spacing(doc, 
        "The following screenshots depict the operational screens captured from the running PakRail system "
        "deployed locally at http://localhost:5173/.", before=6, after=12)

    # Insert screenshots with captions
    # We will loop through the files and add them if they exist.
    screenshots = [
        ('01_home.png', 'Figure 1: PakRail Landing Page showing the primary search interface.'),
        ('15_urdu_home.png', 'Figure 2: Bilingual interface showing PakRail in Urdu with Right-to-Left (RTL) alignment.'),
        ('02_login_page.png', 'Figure 3: Secure Passenger Login interface.'),
        ('03_passenger_dashboard.png', 'Figure 4: Logged-in Passenger Home page displaying customized navbar controls.'),
        ('04_train_search.png', 'Figure 5: Train search interface with pre-selected query inputs.'),
        ('05_search_results.png', 'Figure 6: Train Search results showing list of active trains matching departure and arrival destinations.'),
        ('06_seat_selection.png', 'Figure 7: Real-time visual seat selection grid showing available seats.'),
        ('07_booking_details.png', 'Figure 8: Passenger Details entry page requiring name, age, and 13-digit CNIC.'),
        ('08_payment_simulation.png', 'Figure 9: Checkout step allowing simulated Card payment or Cash/At-Counter hold selection.'),
        ('09_booking_confirmation.png', 'Figure 10: Successful booking confirmation page displaying the generated unique PNR code.'),
        ('10_my_bookings.png', 'Figure 11: Booking history page showing active trips and cancellation options.'),
        ('11_admin_dashboard.png', 'Figure 12: Admin Dashboard rendering aggregated metrics (revenue, active trains, bookings, users).'),
        ('12_admin_train_management.png', 'Figure 13: Admin Train Management panel for adding, updating, or deleting train records.'),
        ('13_admin_booking_management.png', 'Figure 14: Admin Booking Auditing panel for system-wide reservation management.'),
        ('14_admin_reports.png', 'Figure 15: Admin Reports and Analytics displaying visual charts of system load factors and revenue.')
    ]
    
    __dirname = os.path.dirname(os.path.abspath(__file__))
    workspace_root = os.path.abspath(os.path.join(__dirname, '..'))
    local_screenshots_dir = os.path.join(workspace_root, 'screenshots')
    
    for filename, caption in screenshots:
        filepath = os.path.join(local_screenshots_dir, filename)
        if os.path.exists(filepath):
            print(f"Adding image: {filepath}")
            # Center paragraph for image
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(12)
            p_img.paragraph_format.space_after = Pt(6)
            p_img.add_run().add_picture(filepath, width=Inches(5.5))
            
            # Caption paragraph
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(0)
            p_cap.paragraph_format.space_after = Pt(18)
            run_cap = p_cap.add_run(caption)
            run_cap.italic = True
            run_cap.font.size = Pt(9.5)
            run_cap.font.color.rgb = RGBColor(100, 100, 100)
        else:
            print(f"Warning: Screenshot file not found: {filepath}")

    doc.add_page_break()

    # 15. Conclusion
    doc.add_heading("15. Conclusion", level=1)
    p = add_paragraph_with_spacing(doc, 
        "The PakRail project represents a practical, comprehensive digital solution for railway reservation and seat "
        "management. By replacing legacy manual ledger procedures with modern client-server transactional scripting, the "
        "application eliminates double-booking vulnerability, reduces queuing delays, and guarantees record transparency.", 
        before=6, after=6)
    p = add_paragraph_with_spacing(doc, 
        "From a Software Project Management perspective, the project demonstrates technical viability using the high-performance MERN "
        "stack, meets operational workflows of passengers and railway staff, fits within academic schedule boundaries, and "
        "presents an economically suitable cost structure for semester-based environments. Ultimately, PakRail establishes "
        "a reliable template for digitization and user accessibility in national transport infrastructures.", 
        before=6, after=12)

    # 16. Final Project Information
    doc.add_heading("16. Final Project Information", level=1)
    table_fi = doc.add_table(rows=5, cols=2)
    table_fi.alignment = WD_TABLE_ALIGNMENT.CENTER
    fi_data = [
        ['Team Members & SAP IDs', 'Alizay Aman Niazi (58604)\nAyesha Kanwal (57505)'],
        ['Course & University', 'Software Project Management (SPM)\nRiphah International University'],
        ['Academic Instructor', 'Sir Waqar / Dr. Waqar Mehmood'],
        ['Project Title', 'PakRail – Smart Railway Reservation & Seat Management System'],
        ['Application Local URL', 'http://localhost:5173/']
    ]
    for idx, row in enumerate(fi_data):
        cell_lbl = table_fi.rows[idx].cells[0]
        cell_lbl.text = row[0]
        set_cell_background(cell_lbl, HEX_PRIMARY)
        cell_lbl.paragraphs[0].runs[0].font.bold = True
        cell_lbl.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell_lbl, top=100, bottom=100, left=150, right=150)
        
        cell_val = table_fi.rows[idx].cells[1]
        cell_val.text = row[1]
        set_cell_margins(cell_val, top=100, bottom=100, left=150, right=150)
        if idx % 2 != 0:
            set_cell_background(cell_val, HEX_LIGHT_GREY)

    # Save Document
    report_path = os.path.join(workspace_root, 'PakRail_PBL_Report.docx')
    doc.save(report_path)
    console_log = f"SUCCESS: Report saved to {report_path}"
    print(console_log)

if __name__ == '__main__':
    main()
